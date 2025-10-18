"""
Export service for Whysper documentation.

This service provides functionality for exporting generated documentation
to various formats including PDF, HTML, and Word documents.
"""

import os
import io
import tempfile
import subprocess
from typing import Dict, Any, Optional, List
from pathlib import Path
from common.logger import get_logger

logger = get_logger(__name__)


class ExportService:
    """
    Service for exporting documentation to various formats.
    
    This service handles converting markdown documentation to PDF, HTML, and Word
    documents with proper formatting, styling, and structure preservation.
    """
    
    def __init__(self):
        self.logger = get_logger(f"{__name__}.ExportService")
        self.templates_dir = os.path.join(
            os.path.dirname(__file__), 
            "../../templates/export"
        )
    
    def export_to_pdf(self, content: str, options: Dict[str, Any] = None) -> bytes:
        """
        Export documentation to PDF format.
        
        Args:
            content: Markdown content to export
            options: Export options
            
        Returns:
            PDF document as bytes
        """
        self.logger.info("Exporting documentation to PDF")
        
        try:
            options = options or {}
            
            # Convert markdown to HTML
            html_content = self._markdown_to_html(content)
            
            # Apply HTML template
            template_name = options.get('template', 'pdf_template.html')
            html_content = self._apply_html_template(html_content, template_name, options)
            
            # Try to convert HTML to PDF using available tools
            try:
                # Try WeasyPrint if available
                import weasyprint
                pdf_bytes = weasyprint.HTML(string=html_content).write_pdf()
                self.logger.info("PDF export completed with WeasyPrint")
                return pdf_bytes
            except ImportError:
                self.logger.warning("WeasyPrint not available, trying ReportLab")
            
            # Fallback to ReportLab
            pdf_bytes = self._export_to_pdf_with_reportlab(content, options)
            self.logger.info("PDF export completed with ReportLab")
            return pdf_bytes
            
        except Exception as e:
            self.logger.error(f"Error exporting to PDF: {e}")
            # Return simple markdown as fallback
            return content.encode('utf-8')
    
    def export_to_html(self, content: str, options: Dict[str, Any] = None) -> str:
        """
        Export documentation to HTML format.
        
        Args:
            content: Markdown content to export
            options: Export options
            
        Returns:
            HTML document as string
        """
        self.logger.info("Exporting documentation to HTML")
        
        try:
            options = options or {}
            
            # Convert markdown to HTML
            html_content = self._markdown_to_html(content)
            
            # Apply HTML template
            template_name = options.get('template', 'html_template.html')
            html_content = self._apply_html_template(html_content, template_name, options)
            
            self.logger.info("HTML export completed successfully")
            return html_content
            
        except Exception as e:
            self.logger.error(f"Error exporting to HTML: {e}")
            # Return simple HTML as fallback
            return f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Documentation</title>
                <style>
                    body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }}
                    h1, h2, h3 {{ color: #333; }}
                    code {{ background-color: #f5f5f5; padding: 2px 4px; }}
                    pre {{ background-color: #f5f5f5; padding: 10px; overflow-x: auto; }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
    
    def export_to_docx(self, content: str, options: Dict[str, Any] = None) -> bytes:
        """
        Export documentation to Word document format.
        
        Args:
            content: Markdown content to export
            options: Export options
            
        Returns:
            Word document as bytes
        """
        self.logger.info("Exporting documentation to Word document")
        
        try:
            options = options or {}
            
            # Try to use python-docx if available
            try:
                from docx import Document
                from docx.shared import Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                # Create Word document
                doc = Document()
                
                # Add title
                title = doc.add_heading('Documentation', 0)
                
                # Process markdown content
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        doc.add_paragraph()
                    elif line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith('```'):
                        # Code block - skip opening ```
                        continue
                    else:
                        doc.add_paragraph(line)
                
                # Save to buffer
                buffer = io.BytesIO()
                doc.save(buffer)
                
                self.logger.info("Word document export completed successfully")
                return buffer.getvalue()
                
            except ImportError:
                self.logger.warning("python-docx not available")
            
            # Fallback: return markdown as bytes
            self.logger.warning("Returning markdown content as fallback")
            return content.encode('utf-8')
            
        except Exception as e:
            self.logger.error(f"Error exporting to Word document: {e}")
            return content.encode('utf-8')
    
    def _markdown_to_html(self, content: str) -> str:
        """Convert markdown content to HTML."""
        try:
            # Try to use markdown2 if available
            import markdown2
            return markdown2.markdown(content, extras=['tables', 'fenced-code-blocks'])
        except ImportError:
            # Try to use markdown if available
            try:
                import markdown
                return markdown.markdown(content, extensions=['tables', 'fenced_code'])
            except ImportError:
                # Simple fallback conversion
                html = content.replace('\n\n', '</p><p>').replace('\n', '<br>')
                return f'<p>{html}</p>'
    
    def _apply_html_template(self, content: str, template_name: str, options: Dict[str, Any]) -> str:
        """Apply HTML template to content."""
        try:
            template_path = os.path.join(self.templates_dir, template_name)
            
            if os.path.exists(template_path):
                with open(template_path, 'r', encoding='utf-8') as f:
                    template = f.read()
                
                # Replace placeholders
                template = template.replace('{content}', content)
                template = template.replace('{title}', options.get('title', 'Documentation'))
                template = template.replace('{author}', options.get('author', 'Whysper'))
                template = template.replace('{description}', options.get('description', 'Generated documentation'))
                
                return template
            else:
                self.logger.warning(f"Template not found: {template_path}")
                return content
                
        except Exception as e:
            self.logger.error(f"Error applying HTML template: {e}")
            return content
    
    def _export_to_pdf_with_reportlab(self, content: str, options: Dict[str, Any]) -> bytes:
        """Export to PDF using ReportLab as fallback."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            
            # Create buffer
            buffer = io.BytesIO()
            
            # Create PDF document
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            
            # Get styles
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            title_style = styles['Heading1']
            title = Paragraph("Documentation", title_style)
            story.append(title)
            story.append(Spacer(1, 12))
            
            # Process content
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    if line.startswith('# '):
                        p = Paragraph(line[2:], styles['Heading1'])
                    elif line.startswith('## '):
                        p = Paragraph(line[3:], styles['Heading2'])
                    elif line.startswith('### '):
                        p = Paragraph(line[4:], styles['Heading3'])
                    else:
                        p = Paragraph(line, styles['Normal'])
                    
                    story.append(p)
                    story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            
            # Get PDF bytes
            buffer.seek(0)
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            return pdf_bytes
            
        except ImportError:
            self.logger.warning("ReportLab not available")
            return content.encode('utf-8')
        except Exception as e:
            self.logger.error(f"Error with ReportLab PDF export: {e}")
            return content.encode('utf-8')
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported export formats."""
        return ['markdown', 'html', 'pdf', 'docx']
    
    def get_format_options(self, format: str) -> Dict[str, Any]:
        """Get available options for a specific format."""
        options = {
            'pdf': {
                'page_size': ['letter', 'a4'],
                'margin': ['normal', 'narrow', 'wide'],
                'template': ['default', 'minimal', 'professional']
            },
            'html': {
                'theme': ['light', 'dark'],
                'include_toc': [True, False],
                'template': ['default', 'modern', 'minimal']
            },
            'docx': {
                'font_size': [10, 11, 12],
                'template': ['default', 'professional']
            }
        }
        
        return options.get(format, {})


# Singleton instance
export_service = ExportService()